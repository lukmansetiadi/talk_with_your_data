import streamlit as st
import pandas as pd
import os
import docx
import re
from io import BytesIO
from ai_sql_agent import generate_and_execute_sql, split_queries_with_comments
from database_client import execute_sql_query
from schema_exporter import export_schema_to_markdown
from chart_generator import generate_chart_code, execute_chart_code
from gemini_client import request_gemini_llm
from groq_client import request_groq_llm
from bedrock_client import request_bedrock_llm
from prompt_template import (
    SCHEMA_PROMPT_TEMPLATE,
    RULES_PROMPT_TEMPLATE,
    RESEARCH_ANALYSIS_PROMPT_TEMPLATE,
    OUTPUT_ANALYSIS_PROMPT_TEMPLATE,
    QUERY_RECOMMENDATION_PROMPT_TEMPLATE
)
import time

st.set_page_config(page_title="AI SQL Database Agent", layout="wide")

# Use session state to handle the expander's expanded state
if 'edit_schema_expanded' not in st.session_state:
    st.session_state.edit_schema_expanded = False

if 'edit_research_expanded' not in st.session_state:
    st.session_state.edit_research_expanded = False

if 'edit_rules_expanded' not in st.session_state:
    st.session_state.edit_rules_expanded = False

if 'edit_template_expanded' not in st.session_state:
    st.session_state.edit_template_expanded = False

if 'edit_output_expanded' not in st.session_state:
    st.session_state.edit_output_expanded = False

if 'edit_recommendation_expanded' not in st.session_state:
    st.session_state.edit_recommendation_expanded = False

# Prompt session states
if "research_prompt_input_key" not in st.session_state:
    st.session_state.research_prompt_input_key = ""

if "output_research_prompt_input_key" not in st.session_state:
    st.session_state.output_research_prompt_input_key = ""

if "recommendation_prompt_input_key" not in st.session_state:
    st.session_state.recommendation_prompt_input_key = ""

# Functions to update prompts from templates
def fill_research_prompt():
    st.session_state.research_prompt_input_key = RESEARCH_ANALYSIS_PROMPT_TEMPLATE

def fill_output_prompt():
    st.session_state.output_research_prompt_input_key = OUTPUT_ANALYSIS_PROMPT_TEMPLATE
    
def fill_recommendation_prompt():
    st.session_state.recommendation_prompt_input_key = QUERY_RECOMMENDATION_PROMPT_TEMPLATE


# Function to request LLM directly for deep research
def analyze_with_llm(prompt, provider):
    try:
        if provider.lower() == 'groq':
            return request_groq_llm(prompt)
        elif provider.lower() == 'bedrock':
            return request_bedrock_llm(prompt)
        else:
            return request_gemini_llm(prompt)
    except Exception as e:
        return f"Error analyzing data: {e}"


# Function to handle saving and collapsing schema
def save_schema_and_collapse(file_path, content):
    try:
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(content)
        st.session_state.edit_schema_expanded = False
        st.toast("✅ Schema file updated successfully!")
    except Exception as e:
        st.toast(f"❌ Error saving file: {e}")


# Function to handle saving and collapsing business rules
def save_rules_and_collapse(file_path, content):
    try:
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(content)
        st.session_state.edit_rules_expanded = False
        st.toast("✅ Business rules updated successfully!")
    except Exception as e:
        st.toast(f"❌ Error saving file: {e}")


# Function to handle saving and collapsing research
def save_research_and_collapse(file_path, content):
    try:
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(content)
        st.session_state.edit_research_expanded = False
        st.toast("✅ Research file updated successfully!")
    except Exception as e:
        st.toast(f"❌ Error saving file: {e}")


# Function to handle saving and collapsing template
def save_template_and_collapse(file_path, content):
    try:
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(content)
        st.session_state.edit_template_expanded = False
        st.toast("✅ Template file updated successfully!")
    except Exception as e:
        st.toast(f"❌ Error saving file: {e}")


# Function to handle saving and collapsing output
def save_output_and_collapse(file_path, content):
    try:
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(content)
        st.session_state.edit_output_expanded = False
        st.toast("✅ Output file updated successfully!")
    except Exception as e:
        st.toast(f"❌ Error saving file: {e}")


def save_recommendation_and_collapse(file_path, content):
    try:
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(content)
        st.session_state.edit_recommendation_expanded = False
        st.toast("✅ Recommendation file updated successfully!")
    except Exception as e:
        st.toast(f"❌ Error saving file: {e}")


# Function to handle adding SQL query results to research
def save_sql_to_research(question, df, provider):
    research_file = "research_context.md"
    try:
        # DO NOT format the dataframe here, as formatting converts numbers to strings (e.g. "190.000")
        # which breaks downstream LLM analysis that expects raw numbers.
        # Instead, we write the raw numeric dataframe to markdown.
        md_table = df.to_markdown(index=False)

        entry = f"""
## Research Entry (SQL Query)
**Provider:** {provider}
**Question:** {question}

**Data Result:**
{md_table}

---
"""
        with open(research_file, "a", encoding="utf-8") as f:
            f.write(entry)
        st.toast("✅ Added SQL result to Research Context!")
        time.sleep(0.5)
        st.rerun()
    except Exception as e:
        st.toast(f"❌ Error saving research: {e}")


# Function to handle adding prompt and SQL to research template
def save_to_research_template(question, sql_query, provider):
    template_file = "research_template.md"
    try:
        entry = f"""
## Research Template Entry
**Provider:** {provider}
**Question:** {question}

**SQL Query:**
```sql
{sql_query}
```

---
"""
        with open(template_file, "a", encoding="utf-8") as f:
            f.write(entry)
        st.toast("✅ Added to Research Template!")
    except Exception as e:
        st.toast(f"❌ Error saving template: {e}")


# Function to run all queries from research template
def run_all_template_queries(file_path):
    if not os.path.exists(file_path):
        st.warning("Template file not found.")
        return

    output_file = "research_template_output.md"
    # Reset or initialize the output file
    try:
        import time
        with open(output_file, "w", encoding="utf-8") as f:
            f.write(f"# Research Template Execution Results\nGenerated on: {time.strftime('%Y-%m-%d %H:%M:%S')}\n\n---\n")
    except Exception as e:
        st.error(f"Error initializing output file: {e}")
        return

    try:
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        # Regex to find Questions and SQL Queries
        pattern = r"\*\*Question:\*\*\s*(.*?)\n\n\*\*SQL Query:\*\*\n```sql\s*(.*?)\s*```"
        matches = re.findall(pattern, content, re.DOTALL)

        if not matches:
            st.info("No queries found in the template.")
            return

        st.toast(f"🚀 Running {len(matches)} queries from template...")
        
        for question, query in matches:
            question = question.strip()
            query = query.strip()
            # We add a user message for the question first
            st.session_state.messages.append({"role": "user", "content": f"[Template Run] {question}"})
            
            try:
                results = execute_sql_query(query)
                st.session_state.messages.append({
                    "role": "assistant", 
                    "query": query, 
                    "results": results, 
                    "provider": "Template"
                })

                # Save to output file
                if results:
                    df = pd.DataFrame(results)
                    md_table = df.to_markdown(index=False)
                    output_entry = f"## Question: {question}\n\n**SQL Query:**\n```sql\n{query}\n```\n\n**Results:**\n{md_table}\n\n---\n"
                else:
                    output_entry = f"## Question: {question}\n\n**SQL Query:**\n```sql\n{query}\n```\n\n**Results:** No rows returned.\n\n---\n"
                
                with open(output_file, "a", encoding="utf-8") as f:
                    f.write(output_entry)

            except Exception as e:
                error_msg = str(e)
                st.session_state.messages.append({
                    "role": "assistant", 
                    "query": query, 
                    "error": error_msg, 
                    "provider": "Template"
                })
                with open(output_file, "a", encoding="utf-8") as f:
                    f.write(f"## Question: {question}\n\n**SQL Query:**\n```sql\n{query}\n```\n\n**Error:** {error_msg}\n\n---\n")
        
        st.rerun()

    except Exception as e:
        st.error(f"Error running template queries: {e}")

# Function to run all queries from query recommendation file
def run_all_recommendation_queries(file_path):
    if not os.path.exists(file_path):
        st.warning("Recommendation file not found.")
        return
        
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        # Try to parse exact same format as Template first
        pattern = r"\*\*Question:\*\*\s*(.*?)\n\n\*\*SQL Query:\*\*\n```sql\s*(.*?)\s*```"
        matches = re.findall(pattern, content, re.DOTALL)

        if not matches:
             # Fallback to older format parsing if template format not found
             pattern_fallback = r"\*\*Prompt:\*\*\s*(.*?)\n\n\*\*Recommendation:\*\*.*?```sql\s*(.*?)\s*```"
             matches = re.findall(pattern_fallback, content, re.DOTALL)
             
             if not matches:
                 fallback_only_sql = r"```sql\s*(.*?)\s*```"
                 sql_blocks = re.findall(fallback_only_sql, content, re.DOTALL)
                 if sql_blocks:
                      matches = [(f"Executing Query {i+1}", q) for i, q in enumerate(sql_blocks)]
                 else:
                    st.info("No queries found in the recommendation file.")
                    return

        st.toast(f"🚀 Running {len(matches)} queries from recommendations...")
        
        for question, query in matches:
            question = question.strip()
            query = query.strip()
            
            # If the query itself contains semicolons (multiple queries in one block),
            # split them up so they execute properly and output nicely to chat
            sub_queries = split_queries_with_comments(query)
            
            if sub_queries:
                 for sub_desc, sub_q in sub_queries:
                     # Add a user message for each sub-query
                     display_question = f"[Recommendation Run] {question}"
                     if sub_desc and not sub_desc.startswith("Generated Query"):
                         display_question += f"\n\n*(Sub-task: {sub_desc})*"
                         
                     st.session_state.messages.append({"role": "user", "content": display_question})
                     
                     try:
                         results = execute_sql_query(sub_q)
                         st.session_state.messages.append({
                             "role": "assistant", 
                             "query": sub_q, 
                             "results": results, 
                             "provider": "Recommendation"
                         })
     
                     except Exception as e:
                         error_msg = str(e)
                         st.session_state.messages.append({
                             "role": "assistant", 
                             "query": sub_q, 
                             "error": error_msg, 
                             "provider": "Recommendation"
                         })
            else:
                 st.session_state.messages.append({"role": "user", "content": f"[Recommendation Run] {question}"})
                 try:
                      results = execute_sql_query(query)
                      st.session_state.messages.append({
                          "role": "assistant", 
                          "query": query, 
                          "results": results, 
                          "provider": "Recommendation"
                      })
                 except Exception as e:
                      st.session_state.messages.append({
                          "role": "assistant", 
                          "query": query, 
                          "error": str(e), 
                          "provider": "Recommendation"
                      })
        st.rerun()

    except Exception as e:
        st.error(f"Error running recommendation queries: {e}")


# Function to handle adding analysis results to research
def append_analysis_to_research(question, analysis, provider, target_file="research_context.md"):
    try:
        entry = f"""
## Research Entry (Deep Analysis)
**Provider:** {provider}
**Analysis Question:** {question}

**Analysis Result:**
{analysis}

---
"""
        with open(target_file, "a", encoding="utf-8") as f:
            f.write(entry)
        st.toast(f"✅ Added analysis to {target_file}!")
        time.sleep(0.5)
        st.rerun()
    except Exception as e:
        st.toast(f"❌ Error saving research: {e}")


def append_to_recommendation_file(question, recommendation, provider):
    recommendation_file = "query_recommendation.md"
    try:
        entry = f"""
## Query Recommendation
**Provider:** {provider}
**Question:** {question}

**SQL Query:**
```sql
{recommendation}
```

---
"""
        with open(recommendation_file, "a", encoding="utf-8") as f:
            f.write(entry)
        st.toast(f"✅ Added recommendation to {recommendation_file}!")
        time.sleep(0.5)
    except Exception as e:
        st.toast(f"❌ Error saving recommendation: {e}")


# Function to extract text from a .docx file and append to research
def append_docx_to_research(uploaded_file):
    research_file = "research_context.md"
    try:
        doc = docx.Document(uploaded_file)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        doc_content = "\n".join(full_text)

        entry = f"""
## Research Entry (Uploaded Document)
**Source:** {uploaded_file.name}

**Document Content:**
{doc_content}

---
"""
        with open(research_file, "a", encoding="utf-8") as f:
            f.write(entry)
        st.toast("✅ Document content added to Research Context!")
        time.sleep(0.5)
        st.rerun()
    except Exception as e:
        st.toast(f"❌ Error reading/saving document: {e}")


# Helper function to convert dataframe to excel bytes
@st.cache_data(show_spinner=False)
def convert_df_to_excel(df):
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Sheet1')
    processed_data = output.getvalue()
    return processed_data


# Helper function to parse LLM analysis text back into a dataframe (if it contains a markdown table)
def _extract_markdown_table_to_df(markdown_text):
    import io
    try:
        # Find lines that look like a markdown table
        lines = markdown_text.split('\n')
        table_lines = [line for line in lines if '|' in line]

        if not table_lines or len(table_lines) < 3:
            return None

        # Extract the table portion and try to read it with pandas
        table_text = '\n'.join(table_lines)
        # Using pipe as separator, skipping the separator row (which usually contains ---)
        # We read all lines that have pipes, but we might need to filter the row with hyphens
        clean_lines = [line for line in table_lines if not set(line.strip().replace('|', '').replace(' ', '')) == {'-'}]
        clean_text = '\n'.join(clean_lines)

        # Read the pseudo-csv string
        df = pd.read_csv(io.StringIO(clean_text), sep='|', skipinitialspace=True).dropna(axis=1, how='all')
        df.columns = df.columns.str.strip()
        # Clean string values in the dataframe
        for col in df.columns:
            if df[col].dtype == 'object':
                df[col] = df[col].str.strip()

        return df
    except Exception:
        return None


# Sidebar for configuration
with st.sidebar:
    st.header("Configuration")
    selected_llm = st.selectbox(
        "Select LLM Provider",
        options=["groq", "gemini", "bedrock"],
        index=0,
        help="Choose which AI model to use for translating text to SQL."
    )

    # 1. Database Context
    with st.expander("📂 Database Context", expanded=False):
        schema_file = "database_schema.md"
        if os.path.exists(schema_file):
            st.success(f"✅ Schema file `{schema_file}` found.")
            with open(schema_file, "r", encoding="utf-8") as f:
                schema_content = f.read()
            st.download_button("Download Schema MD", schema_content, schema_file, "text/markdown", use_container_width=True)
            if st.button("Toggle Edit Schema", use_container_width=True):
                st.session_state.edit_schema_expanded = not st.session_state.edit_schema_expanded
            if st.session_state.edit_schema_expanded:
                edited_schema = st.text_area("Schema Content", value=schema_content, height=300)
                if st.button("Save Changes", key="save_schema_existing"):
                    save_schema_and_collapse(schema_file, edited_schema)
                    st.rerun()
        else:
            st.warning(f"⚠️ Schema file not found.")
            if st.button("Toggle Create Schema", use_container_width=True):
                st.session_state.edit_schema_expanded = not st.session_state.edit_schema_expanded
            if st.session_state.edit_schema_expanded:
                st.button("Fill Default Template", on_click=lambda: st.session_state.update(schema_content_new=SCHEMA_PROMPT_TEMPLATE))
                edited_schema = st.text_area("Schema Content", value=st.session_state.get("schema_content_new", SCHEMA_PROMPT_TEMPLATE), height=300, key="schema_content_new")
                if st.button("Save Changes", key="save_schema_new"):
                    save_schema_and_collapse(schema_file, edited_schema)
                    st.rerun()
        if st.button("Generate/Refresh Schema from DB", use_container_width=True):
            with st.spinner("Connecting to database..."):
                if export_schema_to_markdown(schema_file):
                    st.toast("✅ Schema successfully updated!")
                    st.session_state.edit_schema_expanded = False
                    time.sleep(1);
                    st.rerun()
                else:
                    st.toast("❌ Failed to update schema.")

    # 2. Business Rules & Terminology
    with st.expander("⚖️ Business Rules & Terminology", expanded=False):
        st.markdown("Inject custom definitions and \"Golden Queries\" to guide the AI.")
        rules_file = "business_rules.md"
        if os.path.exists(rules_file):
            st.success(f"✅ Rules file `{rules_file}` found.")
            with open(rules_file, "r", encoding="utf-8") as f:
                rules_content = f.read()

            if st.button("Toggle Edit Rules", use_container_width=True):
                st.session_state.edit_rules_expanded = not st.session_state.edit_rules_expanded
            if st.session_state.edit_rules_expanded:
                edited_rules = st.text_area("Rules Content", value=rules_content, height=300)
                if st.button("Save Changes", key="save_rules_existing"):
                    save_rules_and_collapse(rules_file, edited_rules)
                    st.rerun()
        else:
            st.warning(f"⚠️ Rules file not found.")
            if st.button("Toggle Create Rules", use_container_width=True):
                st.session_state.edit_rules_expanded = not st.session_state.edit_rules_expanded
            if st.session_state.edit_rules_expanded:
                st.button("Fill Default Template", on_click=lambda: st.session_state.update(rules_content_new=RULES_PROMPT_TEMPLATE))
                edited_rules = st.text_area("Rules Content",
                                            value=st.session_state.get("rules_content_new", RULES_PROMPT_TEMPLATE),
                                            height=300, key="rules_content_new")
                if st.button("Save Changes", key="save_rules_new"):
                    save_rules_and_collapse(rules_file, edited_rules)
                    st.rerun()

    # 3. Research Context
    with st.expander("🔍 Research Context", expanded=False):
        research_file = "research_context.md"

        # Document Upload Section
        uploaded_docx = st.file_uploader("Upload Word Document (.docx)", type=["docx"])
        if uploaded_docx is not None:
            if st.button("Append .docx to Research", use_container_width=True):
                with st.spinner("Extracting text..."):
                    append_docx_to_research(uploaded_docx)

        if os.path.exists(research_file):
            with open(research_file, "r", encoding="utf-8") as f:
                research_content = f.read()
            st.download_button("Download Research MD", research_content, research_file, "text/markdown",
                               use_container_width=True)

            # Toggle the expander state for Research
            if st.button("Toggle Edit Research", use_container_width=True):
                st.session_state.edit_research_expanded = not st.session_state.edit_research_expanded

            if st.session_state.edit_research_expanded:
                edited_research = st.text_area("Research Content", value=research_content, height=300)
                if st.button("Save Changes", key="save_research_existing"):
                    save_research_and_collapse(research_file, edited_research)
                    st.rerun()

            if st.button("Clear Research Data", use_container_width=True):
                os.remove(research_file);
                st.rerun()

            st.markdown("#### Deep Research Analysis")
            
            st.button("Fill Default Template", key="btn_fill_research", on_click=fill_research_prompt)
                 
            research_prompt = st.text_area(
                "Ask the AI to analyze your saved research:",
                placeholder=RESEARCH_ANALYSIS_PROMPT_TEMPLATE,
                height=200,
                value=st.session_state.research_prompt_input_key,
                key="research_prompt_input_key_ui"
            )
            
            if st.button("Run Analysis", use_container_width=True):
                if research_prompt.strip():
                    user_msg = f"**[Deep Research Analysis]**\n{research_prompt}"
                    st.session_state.messages.append({"role": "user", "content": user_msg})
                    with st.spinner(f"Analyzing with {selected_llm.capitalize()}..."):
                        full_prompt = f"Here is the collected research data context:\n\n{research_content}\n\nBased ONLY on the data above, please answer the following request:\n{research_prompt}"
                        analysis_result = analyze_with_llm(full_prompt, selected_llm)
                    # Ensure we set the role, content, and provider keys consistently
                    st.session_state.messages.append(
                        {"role": "assistant", "content": analysis_result, "provider": selected_llm})
                    
                    st.session_state.research_prompt_input_key = ""
                    st.rerun()
                else:
                    st.warning("Please enter a prompt for analysis.")
        else:
            st.info("No research data saved yet.")

    # 4. Research Context Advanced
    with st.expander("🛠️ Research Context Advanced", expanded=False):
        # Sub-section: Research Template
        st.markdown("### 📝 Research Template")
        template_file = "research_template.md"
        if os.path.exists(template_file):
            with open(template_file, "r", encoding="utf-8") as f:
                template_content = f.read()
            st.download_button("Download Template MD", template_content, template_file, "text/markdown",
                               use_container_width=True, key="dl_template")

            if st.button("Toggle Edit Template", use_container_width=True, key="edit_template"):
                st.session_state.edit_template_expanded = not st.session_state.edit_template_expanded

            if st.session_state.edit_template_expanded:
                edited_template = st.text_area("Template Content", value=template_content, height=300, key="ta_template")
                if st.button("Save Changes", key="save_template_existing"):
                    save_template_and_collapse(template_file, edited_template)
                    st.rerun()

            if st.button("🚀 Run All Queries", use_container_width=True, key="run_template"):
                run_all_template_queries(template_file)

            if st.button("Clear Template Data", use_container_width=True, key="clear_template"):
                os.remove(template_file);
                st.rerun()
        else:
            st.info("No research template saved yet.")

        st.markdown("---")
        
        # Sub-section: Research Template Output
        st.markdown("### 📊 Research Template Output")
        output_file = "research_template_output.md"
        if os.path.exists(output_file):
            with open(output_file, "r", encoding="utf-8") as f:
                output_content = f.read()
            st.download_button("Download Output MD", output_content, output_file, "text/markdown",
                               use_container_width=True, key="dl_output")

            if st.button("Toggle Edit Output", use_container_width=True, key="edit_output"):
                st.session_state.edit_output_expanded = not st.session_state.edit_output_expanded

            if st.session_state.edit_output_expanded:
                edited_output = st.text_area("Output Content", value=output_content, height=300, key="ta_output")
                if st.button("Save Changes", key="save_output_existing"):
                    save_output_and_collapse(output_file, edited_output)
                    st.rerun()

            if st.button("Clear Output Data", use_container_width=True, key="clear_output"):
                os.remove(output_file);
                st.rerun()

            st.markdown("#### Deep Research Analysis (Template Output)")
            
            # Sub-section: Upload document for additional context
            uploaded_output_docx = st.file_uploader("Upload Word Document (.docx) for additional context", type=["docx"], key="output_docx_uploader")
            doc_context = ""
            if uploaded_output_docx is not None:
                try:
                    doc = docx.Document(uploaded_output_docx)
                    doc_context = "\n".join([para.text for para in doc.paragraphs])
                    st.success(f"✅ Context from `{uploaded_output_docx.name}` loaded.")
                except Exception as e:
                    st.error(f"Error reading document: {e}")
            
            st.button("Fill Default Template", key="btn_fill_output", on_click=fill_output_prompt)
                
            output_research_prompt = st.text_area(
                "Ask the AI to analyze your batch results (optionally with uploaded doc):",
                placeholder=OUTPUT_ANALYSIS_PROMPT_TEMPLATE,
                height=200,
                value=st.session_state.output_research_prompt_input_key,
                key="output_research_prompt_input_key_ui"
            )
            
            if st.button("Run Analysis", use_container_width=True, key="run_output_analysis"):
                if output_research_prompt.strip():
                    user_msg = f"**[Deep Research Analysis: Template Output]**\n{output_research_prompt}"
                    if uploaded_output_docx:
                        user_msg += f"\n*(Context includes: {uploaded_output_docx.name})*"
                    
                    st.session_state.messages.append({"role": "user", "content": user_msg})
                    with st.spinner(f"Analyzing with {selected_llm.capitalize()}..."):
                        # Build combined context
                        combined_context = f"--- BATCH SQL RESULTS ---\n{output_content}\n\n"
                        if doc_context:
                            combined_context += f"--- ADDITIONAL DOCUMENT CONTEXT ({uploaded_output_docx.name}) ---\n{doc_context}\n\n"
                        
                        full_prompt = f"Here is the context for analysis:\n\n{combined_context}\n\nBased ONLY on the data provided above, please answer the following request:\n{output_research_prompt}"
                        analysis_result = analyze_with_llm(full_prompt, selected_llm)
                    
                    st.session_state.messages.append(
                        {"role": "assistant", "content": analysis_result, "provider": selected_llm})
                    
                    st.session_state.output_research_prompt_input_key = ""
                    st.rerun()
                else:
                    st.warning("Please enter a prompt for analysis.")
        else:
            st.info("No template output generated yet.")

    # 5. Query Recommendations
    with st.expander("💡 Query Recommendations", expanded=False):
        st.markdown("Upload markdown files (e.g., schema, research output) and prompt the AI to recommend new queries.")
        
        uploaded_md_files = st.file_uploader("Upload Context Files (.md)", type=["md"], accept_multiple_files=True, key="recommendation_uploader")
        
        st.button("Fill Default Template", key="btn_fill_recommendation", on_click=fill_recommendation_prompt)
             
        recommendation_prompt = st.text_area(
            "Prompt for Recommendations:",
            placeholder=QUERY_RECOMMENDATION_PROMPT_TEMPLATE,
            height=150,
            value=st.session_state.recommendation_prompt_input_key,
            key="recommendation_prompt_input_key_ui"
        )
        
        if st.button("Generate Query Recommendations", use_container_width=True):
            if recommendation_prompt.strip():
                if not uploaded_md_files:
                     st.warning("Please upload at least one context file.")
                else:
                    context_content = ""
                    for file in uploaded_md_files:
                        try:
                             content = file.getvalue().decode("utf-8")
                             context_content += f"--- CONTENT FROM: {file.name} ---\n{content}\n\n"
                        except Exception as e:
                             st.error(f"Error reading {file.name}: {e}")
                    
                    user_msg = f"**[Query Recommendations]**\n{recommendation_prompt}\n*(Using {len(uploaded_md_files)} uploaded context file(s))*"
                    st.session_state.messages.append({"role": "user", "content": user_msg})
                    
                    with st.spinner(f"Generating recommendations with {selected_llm.capitalize()}..."):
                        full_prompt = f"Here is the context for analysis:\n\n{context_content}\n\nBased on the data provided above, please answer the following request:\n{recommendation_prompt}"
                        recommendation_result = analyze_with_llm(full_prompt, selected_llm)
                        
                        # Save the recommendation
                        append_to_recommendation_file(recommendation_prompt, recommendation_result, selected_llm)

                    st.session_state.messages.append(
                        {"role": "assistant", "content": recommendation_result, "provider": selected_llm})
                        
                    st.session_state.recommendation_prompt_input_key = ""
                    st.rerun()
            else:
                st.warning("Please enter a prompt.")

        st.markdown("---")
        st.markdown("### 📝 Manage Query Recommendations")
        recommendation_file = "query_recommendation.md"
        if os.path.exists(recommendation_file):
            with open(recommendation_file, "r", encoding="utf-8") as f:
                recommendation_content = f.read()
            st.download_button("Download Recommendation MD", recommendation_content, recommendation_file, "text/markdown",
                               use_container_width=True, key="dl_recommendation")

            if st.button("Toggle Edit Query Recommendation", use_container_width=True, key="edit_recommendation"):
                st.session_state.edit_recommendation_expanded = not st.session_state.edit_recommendation_expanded

            if st.session_state.edit_recommendation_expanded:
                edited_recommendation = st.text_area("Recommendation Content", value=recommendation_content, height=300, key="ta_recommendation")
                if st.button("Save Changes", key="save_recommendation_existing"):
                    save_recommendation_and_collapse(recommendation_file, edited_recommendation)
                    st.rerun()
                    
            if st.button("🚀 Run All Recommendation Queries", use_container_width=True, key="run_recommendation_queries"):
                run_all_recommendation_queries(recommendation_file)

            if st.button("Clear Recommendation Data", use_container_width=True, key="clear_recommendation"):
                os.remove(recommendation_file)
                st.rerun()
        else:
            st.info("No recommendations saved yet.")

st.title("🤖 Text-to-SQL AI Chatbot")


def format_dataframe_for_display(df):
    """
    Formats numerical columns in a DataFrame to use thousands separators for UI display.
    """
    formatted_df = df.copy()
    for col in formatted_df.columns:
        if pd.api.types.is_numeric_dtype(formatted_df[col]):
            formatted_df[col] = formatted_df[col].apply(lambda x: f"{x:,.0f}".replace(",", ".") if pd.notnull(x) else x)
    return formatted_df


# Clear chat button
col1, col2 = st.columns([8, 1])
with col2:
    if st.button("Clear Chat"):
        st.session_state.messages = []
        st.rerun()

if "messages" not in st.session_state:
    st.session_state.messages = []

for i, message in enumerate(st.session_state.messages):
    if "role" not in message:
        continue

    with st.chat_message(message["role"]):
        if "content" in message:
            st.markdown(message["content"])

        if message["role"] == "assistant":
            if "query" in message and message["query"]:
                with st.expander(f"View Generated SQL Query ({message.get('provider', 'unknown')})", expanded=False):
                    st.code(message["query"], language="sql")
            if "results" in message and message["results"] is not None:
                if len(message["results"]) > 0:

                    df = pd.DataFrame(message["results"])
                    for col in df.columns:
                        # Attempt to convert object columns to numeric, coercing errors
                        if df[col].dtype == 'object':
                            df[col] = pd.to_numeric(df[col], errors='coerce').fillna(df[col])

                    st.dataframe(format_dataframe_for_display(df))

                    # Create columns for the action buttons to sit side-by-side
                    btn_col1, btn_col2, btn_col3, btn_col4 = st.columns([2, 2, 2, 4])

                    with btn_col1:
                        if st.button("📊 Generate Chart", key=f"gen_chart_{i}"):
                            with st.spinner("AI is generating a chart..."):
                                user_question = st.session_state.messages[i - 1].get("content", "Unknown Question")
                                chart_code = generate_chart_code(df, user_question, selected_llm)
                                if chart_code:
                                    st.session_state.messages[i]['chart_code'] = chart_code
                                    st.rerun()
                                else:
                                    st.error("Could not generate chart code.")

                    with btn_col2:
                        if st.button("💾 Save to Research", key=f"save_sql_{i}"):
                            user_question = st.session_state.messages[i - 1].get("content", "Unknown Question")
                            # Save the raw, unformatted dataframe to research
                            save_sql_to_research(user_question, df, message.get('provider', 'unknown'))

                    with btn_col3:
                        if st.button("📝 Save to Template", key=f"save_template_{i}"):
                            user_question = st.session_state.messages[i - 1].get("content", "Unknown Question")
                            save_to_research_template(user_question, message.get("query", ""),
                                                       message.get('provider', 'unknown'))

                    with btn_col4:
                        # Generate the excel file in memory
                        excel_data = convert_df_to_excel(df)
                        # We use the index 'i' in the filename to ensure uniqueness if they download multiple
                        file_name = f"query_results_{i}.xlsx"
                        st.download_button(
                            label="📥 Download Excel",
                            data=excel_data,
                            file_name=file_name,
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            key=f"download_excel_{i}"
                        )
                else:
                    st.info("Query executed successfully, but returned 0 rows.")

            if 'chart_code' in message and message['chart_code']:
                with st.expander("View Generated Chart", expanded=True):
                    raw_df = pd.DataFrame(message['results'])
                    for col in raw_df.columns:
                        if raw_df[col].dtype == 'object':
                            raw_df[col] = pd.to_numeric(raw_df[col], errors='coerce').fillna(raw_df[col])
                    fig, error = execute_chart_code(message['chart_code'], raw_df)
                    if error:
                        st.error(f"Error executing chart code: {error}")
                        st.code(message['chart_code'], language='python')
                    else:
                        st.plotly_chart(fig, use_container_width=True)

            elif "content" in message and not "query" in message and "error" not in message:
                # Create columns for the deep research action buttons
                btn_col1, btn_col2 = st.columns([2, 8])

                with btn_col1:
                    if st.button("💾 Save Analysis", key=f"save_analysis_{i}"):
                        user_question = st.session_state.messages[i - 1].get("content", "Unknown Question")
                        # Detect if this was a template analysis by checking the prefix we added
                        if "**[Deep Research Analysis: Template Output]**" in user_question:
                            target = "research_template_output.md"
                        else:
                            target = "research_context.md"
                        
                        append_analysis_to_research(user_question, message["content"],
                                                    message.get('provider', 'unknown'),
                                                    target_file=target)

                # Check if the analysis text generated by the LLM contains a markdown table
                # If it does, we can extract it to allow downloading as excel
                if "|" in message["content"] and "\n-" in message["content"]:
                    extracted_df = _extract_markdown_table_to_df(message["content"])
                    if extracted_df is not None and len(extracted_df) > 0:
                        with btn_col2:
                            excel_data = convert_df_to_excel(extracted_df)
                            st.download_button(
                                label="📥 Download Embedded Table to Excel",
                                data=excel_data,
                                file_name=f"analysis_table_{i}.xlsx",
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                key=f"download_analysis_excel_{i}"
                            )

            if "error" in message:
                st.error(message["error"])

if prompt := st.chat_input("E.g., Show me the top 10 products that sell the most"):
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)
    with st.chat_message("assistant"):
        # If the user input contains a semicolon, we can assume it's a multi-query input
        # Note: If it's a simple text input it shouldn't contain a semicolon. 
        # But just in case, we check if they pasted a large raw query block
        if ";" in prompt and "```sql" in prompt:
             sub_queries = split_queries_with_comments(prompt)
             for sub_desc, sub_q in sub_queries:
                 # Add a user message for each sub-query
                 st.session_state.messages.append({"role": "user", "content": f"Executing: {sub_desc}"})
                 
                 # Execute each query
                 sql_query, results = generate_and_execute_sql(sub_q, max_retries=3, llm_provider=selected_llm)
                 response_message = {"role": "assistant", "provider": selected_llm}
                 if sql_query:
                     response_message["query"] = sql_query
                     if results is not None:
                         response_message["results"] = results
                     else:
                         response_message["error"] = "Failed to execute query."
                 else:
                     response_message["error"] = f"Failed to generate a valid SQL query."
                 st.session_state.messages.append(response_message)
        else:
            with st.spinner(f'Translating to SQL with {selected_llm.capitalize()}...'):
                sql_query, results = generate_and_execute_sql(prompt, max_retries=3, llm_provider=selected_llm)
                
                # Try to handle the case where the LLM generated multiple queries
                if sql_query and ";" in sql_query:
                     sub_queries = split_queries_with_comments(sql_query)
                     if len(sub_queries) > 1:
                         # We had multiple queries generated. We will display them as separate messages.
                         for i, (sub_desc, sub_q) in enumerate(sub_queries):
                              # We don't want to re-generate the SQL, just execute what was already generated
                              # We only append a new user message for the sub-queries after the first one 
                              # because the very first query corresponds directly to the prompt already shown
                              if i > 0:
                                  st.session_state.messages.append({"role": "user", "content": f"*(Executing Sub-task: {sub_desc})*"})
                              try:
                                  sub_results = execute_sql_query(sub_q)
                                  st.session_state.messages.append({
                                      "role": "assistant",
                                      "query": sub_q,
                                      "results": sub_results,
                                      "provider": selected_llm
                                  })
                              except Exception as e:
                                  st.session_state.messages.append({
                                      "role": "assistant",
                                      "query": sub_q,
                                      "error": str(e),
                                      "provider": selected_llm
                                  })
                         st.rerun()
                         
                # Standard single query handling if split_queries didn't trigger
                response_message = {"role": "assistant", "provider": selected_llm}
                if sql_query:
                    response_message["query"] = sql_query
                    if results is not None:
                        response_message["results"] = results
                    else:
                        response_message["error"] = "Failed to execute query."
                else:
                    response_message["error"] = f"Failed to generate a valid SQL query."
                st.session_state.messages.append(response_message)
        st.rerun()
